"""
gen_final_set.py - Builds ONE combined Word document containing the
ENTIRE AGM final set, in the same order as Final_Set__With_Fill_Data_.pdf:

  1. Notice of AGM (with Notes, Explanatory Statement)
  2. Director's Report
  3. Form AOC-2
  4. Equity List of Shareholders          [tabular page]
  5. Attendance Sheet of Members in AGM   [tabular page]
  6. List of Directors                    [tabular page]
  7. Director's Attendance Sheet          [tabular page]
  8. Resolution - Director Regularization
  9. Resolution - Related Party
 10. Resolution - Auditor Reappointment
 11. Intimation Letter to Auditor

The whole set comes out as a single .docx file:
    /home/claude/agm_automation/output/Final_Set.docx
"""

import os
import re
import sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT, WD_SECTION

from loader import load_master, yes, fmt_num
from docx_helpers import (
    new_doc, add_letterhead, add_run, add_para, add_centered_heading,
    add_section_heading, add_bordered_table, add_keyvalue_table,
    add_aoc2_block_table, add_aoc2_count_box,
    add_two_signer_block, add_single_signer_block,
    _shade_cell, _set_cell_borders,
)


# =====================================================================
# Helpers
# =====================================================================
def _justify(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text)
    return p


def _heading_numbered(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    add_run(p, f"{num}. ", bold=True)
    add_run(p, text, bold=True, underline=True)
    return p


def _heading_lettered(doc, letter, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    add_run(p, f"{letter}) ", bold=True)
    add_run(p, text, bold=True, underline=True)
    return p


_PERSON_HONORIFICS = ("mr.", "mrs.", "ms.", "miss", "dr.", "shri", "smt.", "smt", "shri.")


def _ms_prefix(name):
    """M/s prefix removed globally per user preference. Always returns ''."""
    return ""


def _rpt_titled(rpt):
    """'Title Name' if the user picked a Title (M/s / Mr. / Mrs. / Dr. / …) on
    the RelatedParty tab AND the typed Name doesn't already start with that
    honorific. Avoids 'Mr. Mr. AK Sen' / 'M/s M/S ABC' duplicates."""
    t = str(rpt.get("title", "") or "").strip()
    n = str(rpt.get("name", "") or "").strip()
    if not t:
        return n
    nl = n.lower()
    # Any honorific already at the start of Name? Don't prepend Title.
    for h in ("mr.", "mrs.", "ms.", "miss", "dr.", "shri", "smt.", "smt",
              "m/s ", "m/s.", "m/s"):
        if nl.startswith(h):
            return n
    return f"{t} {n}".strip()


def _currency_text(raw):
    """Turn a max-value cell into 'Rs X,XX,XXX/- (Rupees X Only)'. Understands
    Indian unit suffixes — '10 Crore', '5 Lakhs', '50 Thousand', '500' — and
    multiplies accordingly. Falls back to the raw string if unparseable."""
    if raw is None or str(raw).strip() == "":
        return ""
    s = str(raw).strip()
    import re as _re
    m = _re.match(r"\s*(?:Rs\.?\s*|INR\s*)?([\d,]+(?:\.\d+)?)\s*([A-Za-z]+)?", s)
    if not m:
        return s
    num_str, unit = m.group(1), (m.group(2) or "").lower().rstrip("s.")
    try:
        n = float(num_str.replace(",", ""))
    except (ValueError, TypeError):
        return s
    unit_mult = {
        "crore":    10_000_000,
        "cr":       10_000_000,
        "lakh":     100_000,
        "lac":      100_000,
        "lakhs":    100_000,
        "thousand": 1_000,
        "k":        1_000,
        "hundred":  100,
    }.get(unit, 1)
    n_int = int(n * unit_mult)
    from share_capital_text import indian_comma, indian_words
    return f"Rs {indian_comma(n_int)}/- (Rupees {indian_words(n_int)} Only)"


def _num_in_words(n):
    nums = {1: "One", 2: "Two", 3: "Three", 4: "Four", 5: "Five",
            6: "Six", 7: "Seven", 8: "Eight", 9: "Nine", 10: "Ten",
            11: "Eleven", 12: "Twelve"}
    return nums.get(n, str(n))


def _format_date_long(date_str):
    """'30/09/2025' -> '30th September, 2025'. Falls back to input on error."""
    if not date_str or "/" not in str(date_str):
        return date_str or ""
    try:
        d, m, y = str(date_str).split("/")
        d_i, m_i, y_i = int(d), int(m), int(y)
    except (ValueError, AttributeError):
        return date_str
    if 11 <= d_i % 100 <= 13:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(d_i % 10, "th")
    months = ["", "January", "February", "March", "April", "May", "June",
              "July", "August", "September", "October", "November", "December"]
    if not 1 <= m_i <= 12:
        return date_str
    return f"{d_i}{suffix} {months[m_i]}, {y_i}"


def _fy_end_year(date_str):
    """'31/03/2025' -> '2025'. Falls back to input on error."""
    if not date_str or "/" not in str(date_str):
        return str(date_str or "")
    parts = str(date_str).split("/")
    return parts[-1] if len(parts) == 3 else str(date_str)


def _new_page_with_letterhead(doc, data, first=False):
    """Insert a page-break (unless first) and re-print the company letterhead."""
    if not first:
        doc.add_page_break()
    add_letterhead(doc, data)


def _set_section_orientation(section, landscape):
    """Apply A4 portrait or landscape dimensions + 0.7/0.8" margins to a section."""
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        # Tighter side margins on landscape pages so wide tables (Director
        # Attendance, Shareholders) get the extra width and dates don't wrap.
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def _begin_landscape_page(doc, data):
    """Start a new landscape page with the company letterhead."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_section_orientation(section, landscape=True)
    add_letterhead(doc, data)


def _begin_portrait_page(doc, data):
    """Start a new portrait page with the company letterhead (after landscape)."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_section_orientation(section, landscape=False)
    add_letterhead(doc, data)


# =====================================================================
# SECTION 1: NOTICE
# =====================================================================
def section_notice(doc, data):
    company = data["company"]
    agm = data["agm"]
    signing = data["signing"]
    toggles = data["toggles"]
    aud = data["auditor"]
    regs = data.get("regularizations", []) if yes(toggles, "RegulariseAdditionalDirector") else []
    rpts = data.get("rpts", []) if yes(toggles, "RPT_Required") else []

    add_centered_heading(doc, "NOTICE", size=13, underline=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "NOTICE ", bold=True)
    add_run(p, "is hereby given that the ")
    add_run(p, f"{agm['number']} ", bold=True)
    add_run(p, "Annual General Meeting of the Members of ")
    add_run(p, f"{company['name']} ", bold=True)
    add_run(p, "will be held on ")
    add_run(p, f"{agm['day'].title()}, {agm['date_words']} ", bold=True)
    add_run(p, "at ")
    add_run(p, f"{agm['time']} ", bold=True)
    add_run(p, "at the Registered office of the company situated at ")
    add_run(p, f"{agm['venue']}", bold=True)
    add_run(p, " to transact the following business:")

    add_section_heading(doc, "ORDINARY BUSINESS:")
    ord_count = 0

    ord_count += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, f"{ord_count}. ", bold=True)
    add_run(p, "To receive, consider and adopt the Audited Balance Sheet as on ")
    add_run(p, agm["fy_end_date"], bold=True)
    add_run(p, ", the Profit and Loss Account for the year ended on that date and the Reports of the Directors and Auditors thereon.")

    if yes(toggles, "AuditorReappointment"):
        ord_count += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, f"{ord_count}. ", bold=True)
        add_run(p, "To re-appoint ")
        add_run(p, f"{aud['firm_name']}, {aud['designation']} (Firm Registration No. {aud['frn']}) ", bold=True)
        add_run(p, "as Statutory Auditor of the Company and pass the following resolution:")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # left_indent removed — keep resolution body aligned with item text above
        add_run(p, "\u201cResolved that ", bold=True)
        add_run(p, "pursuant to the provisions of Section 139 and other applicable provisions, if any, of the Companies Act, 2013 and the Rules framed there under, as amended from time to time, ")
        add_run(p, f"{aud['firm_name']}, {aud['designation']} (Firm Registration No. {aud['frn']}) ", bold=True)
        add_run(p, f"be and are hereby re-appointed as Auditors of the Company to hold office from the conclusion of this Annual General Meeting till the conclusion of the Annual General Meeting held in F.Y. {aud['tenure_end_fy']} of the Company, at such remuneration plus service tax, out of pocket, travelling and living expenses, etc., as may be mutually agreed between the Board of Directors of the Company and the Auditors.\u201d")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # left_indent removed — keep resolution body aligned with item text above
        add_run(p, "\u201cResolved further that ", bold=True)
        add_run(p, "any Director of the company be and are hereby individually/severally authorized to digitally sign and file e-form ADT-1 with the Registrar of Companies and to do all such things, deeds, acts which may deem necessary to give effect of the aforesaid resolution.\u201d")

    has_special = bool(regs) or bool(rpts)
    if has_special:
        add_section_heading(doc, "SPECIAL BUSINESS:")

    sno = ord_count
    for dc in regs:
        sno += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, f"{sno}. ", bold=True)
        add_run(p, "To regularize the appointment of ")
        add_run(p, f"{dc['name']} (DIN: {dc['din']}), ", bold=True)
        add_run(p, "as the Director of the Company.")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # left_indent removed — keep resolution body aligned with item text above
        add_run(p, "\u201cResolved that ", bold=True)
        add_run(p, "in accordance with the provision of Section 161(1), read with Section 152 of the Companies Act, 2013 and the Rules made there under (including any statutory modification(s) or re-enactment thereof), and the Article of Association of the Company, ")
        add_run(p, f"{dc['name']} (DIN: {dc['din']}), ", bold=True)
        add_run(p, "be and is hereby appointed as a Director of the Company who was appointed as an Additional Director of the Company by the Board of Directors at its meeting held on ")
        add_run(p, f"{dc['appt_dotted']}", bold=True)
        add_run(p, ".\u201d")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # left_indent removed — keep resolution body aligned with item text above
        add_run(p, "\u201cResolved Further That ", bold=True)
        add_run(p, "any Director of the Company of the company be and are hereby individually/severally authorized to digitally sign and file e-form DIR-12 with the Registrar of Companies and to do all such things, deeds, acts which may deem necessary to give effect of the aforesaid resolution.\u201d")

    for rpt in rpts:
        sno += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, f"{sno}. ", bold=True)
        add_run(p, "To approve the related party transactions under section 188 of the Companies Act, 2013 with ")
        add_run(p, f"{_rpt_titled(rpt)}.", bold=True)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # left_indent removed — keep resolution body aligned with item text above
        add_run(p, "To consider and if thought fit, to pass, with or without modification(s), the following resolution as Ordinary Resolution:")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # left_indent removed — keep resolution body aligned with item text above
        add_run(p, "\u201cResolved That ", bold=True)
        add_run(p, "pursuant to the provisions of Section 188 of the Companies Act, 2013 (\u201cAct\u201d) and other applicable provisions, if any, read with Rule 15 of the Companies (Meetings of Board and its Powers) Rules, 2014, (including any amendments, modifications, variations or re-enactments thereof for the time being in force), the consent of the members of the Company be and is hereby accorded for carrying out and / or continuing with arrangements and transactions with ")
        add_run(p, f"{_rpt_titled(rpt)} ", bold=True)
        add_run(p, f"related party of the Company with respect to {rpt['nature_of_contract'].lower()}, at arm\u2019s length basis and in the ordinary course of business, at arm\u2019s length basis and in the ordinary course of business, on such terms and conditions as the Board of Directors may deem fit up to a maximum aggregate value of ")
        add_run(p, _currency_text(rpt["max_value"]), bold=True)
        add_run(p, " each.\u201d")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # left_indent removed — keep resolution body aligned with item text above
        add_run(p, "\u201cResolved Further That ", bold=True)
        add_run(p, "for the purpose of giving effect to the above resolution, the Board of Directors of the Company be and are hereby authorized to do all acts, deeds and things in their absolute discretion that may be considered necessary, proper and expedient or incidental for the purpose of giving effect to this resolution in the interest of the Company.\u201d")

    add_single_signer_block(
        doc, signer=signing["first"],
        date_str=agm["notice_date"], place_str=agm["notice_place"],
        closing="By Order of the Board of Directors", company_name=None,
        reg_office_address=company["address"],
    )

    # Notes
    _new_page_with_letterhead(doc, data)
    add_section_heading(doc, "NOTES:")
    notes = [
        "A Member entitled to attend and vote is entitled to appoint a proxy to attend and vote instead of himself and the proxy need not be a member. Proxies in order to be effective must be received by the company not later than forty-eight (48) hours before the meeting. Proxies submitted on behalf of limited companies, societies, etc., must be supported by appropriate resolutions/authority, as applicable. A person can act as a proxy on behalf of Members not exceeding fifty in number and holding in the aggregate not more than ten percent of the total share capital of the Company carrying voting rights. A Member holding more than ten percent of the total share capital of the Company carrying voting rights may appoint a single person as a proxy and such person shall not act as proxy for any other person or shareholder.",
        "In case of joint holders attending the Meeting, only such joint holder who is higher in the order of names will be entitled to vote at the Meeting.",
        "Relevant documents referred to in the accompanying Notice and in the Explanatory Statements are open for inspection by the Members at the Company\u2019s Registered Office on all working days of the Company, during business hours up to the date of the Meeting.",
        "Corporate Members intending to send their authorized representatives to attend the Meeting pursuant to Section 113 of the Companies Act, 2013 are requested to send to the Company, a certified copy of the relevant Board Resolution together with their respective specimen signatures authorizing their representative(s) to attend and vote on their behalf at the Meeting.",
        "Members seeking any information with regard to the Accounts are requested to write to the Company at an early date, so as to enable the Management to keep the information ready at the meeting.",
        "The Notice of AGM, Annual Report, Proxy Form and Attendance Slip are being sent to the Members.",
    ]
    for i, note in enumerate(notes, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, f"{i}. ", bold=True)
        add_run(p, note)

    if has_special:
        _new_page_with_letterhead(doc, data)
        add_centered_heading(
            doc,
            "EXPLANATORY STATEMENT IN RESPECT OF THE SPECIAL BUSINESS PURSUANT TO SECTION 102(1) OF THE COMPANIES ACT, 2013",
            size=11, underline=True,
        )
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, "As required by Section 102(1) of the Companies Act, 2013 (the Act), the following Explanatory Statement set out the material facts relating to the Special Business:")

        item_no = ord_count
        for dc in regs:
            item_no += 1
            p = doc.add_paragraph()
            add_run(p, f"Item No. {item_no}: ", bold=True, underline=True)
            add_run(p, f"To regularize the appointment of {dc['name']} (DIN: {dc['din']}) as the Director of the Company.", bold=True, underline=True)

            # Matches the reference Explanatory Statement wording exactly:
            # "X (DIN:…) was appointed as an Additional Director in the Meeting
            # of Board of Directors held on <date>. The Board feels that
            # presence of X on the Board is desirable and would be beneficial
            # to the company. Thus, Board recommend to the shareholders for
            # the regularization of Director in accordance with the provision
            # of Section 161(1), read with Section 152 of the Companies Act, 2013."
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, f"{dc['name']} (DIN: {dc['din']}) ", bold=True)
            add_run(p, "was appointed as an Additional Director in the Meeting of Board of Directors held on ")
            add_run(p, f"{dc['appt_dotted']}", bold=True)
            add_run(p, f". The Board feels that presence of {dc['name']} on the Board is desirable and would be beneficial to the company. Thus, Board recommend to the shareholders for the regularization of Director in accordance with the provision of Section 161(1), read with Section 152 of the Companies Act, 2013.")

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, f"In the opinion of the Board and the disclosures made by {dc['name']}, he/she fulfills all the criteria for appointment as Director of the Company.")

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, f"Accordingly, the Board recommends passing of the resolution at Item No. {item_no} of the Notice as an Ordinary Resolution.")

        for rpt in rpts:
            item_no += 1
            p = doc.add_paragraph()
            add_run(p, f"Item No. {item_no}: ", bold=True)
            add_run(p, f"To approve the related party transactions under section 188 of the Companies Act, 2013 with {_rpt_titled(rpt)}", bold=True)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, "As per Section 188 of the Companies Act, 2013 (\u201cthe Act\u201d), transactions with related parties which are on arm\u2019s length basis and in the ordinary course of business, approval of shareholders is required for sale or purchase of goods or services, amounting to 10% or more of the turnover of the Company.")

            p = doc.add_paragraph()
            add_run(p, "Details of contracts or arrangements or transactions at Arm\u2019s length basis:", bold=True)

            rpt_table = [
                ["a)", "Name of the related party", _rpt_titled(rpt)],
                ["b)", "Name of the director or key managerial personnel who is related, if any", rpt["related_director"]],
                ["c)", "Nature of relationship", rpt["relationship"]],
                ["d)", "Nature, material terms, monetary value and particulars of the contract or arrangements",
                 f"{rpt['nature_of_contract']}.\nThe monetary value of the transaction up to a maximum aggregate value of {_currency_text(rpt['max_value'])} each."],
                ["e)", "Any other information relevant or important for the members to take a decision on the proposed resolution", "NA"],
            ]
            add_bordered_table(doc, ["S. No.", "Particulars", "Details"], rpt_table,
                               col_widths_inches=[0.7, 3.0, 3.5])

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, f"The Board recommends the Ordinary Resolution as set out in Item No. {item_no} of this Notice for approval of the Members.")

    # Route Map
    _new_page_with_letterhead(doc, data)
    add_centered_heading(doc, "ROUTE MAP", size=13, underline=True)
    add_para(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "(Image of the location at which the Meeting will be held)", italic=True)


# =====================================================================
# SECTION 2: DIRECTOR'S REPORT
# =====================================================================
def section_directors_report(doc, data):
    company = data["company"]
    agm = data["agm"]
    signing = data["signing"]
    fin = data["financials"]
    toggles = data["toggles"]
    emp = data["employees"]
    sc = data["share_capital"]
    aud = data["auditor"]
    regs = data.get("regularizations", []) if yes(toggles, "RegulariseAdditionalDirector") else []
    rpts = data.get("rpts", []) if yes(toggles, "RPT_Required") else []

    add_centered_heading(doc, "DIRECTOR\u2019S REPORT", size=13, underline=True)

    add_para(doc, "")
    add_para(doc, "To,")
    add_para(doc, "The Members,")
    p = doc.add_paragraph()
    add_run(p, company["name"], bold=True)

    add_para(doc, "")
    _justify(doc, f"Your directors have pleasure in presenting their {agm['number']} Annual Report on the business and operations of the Company and the accounts for the Financial Year ended {agm['fy_end_date']}.")

    # 1. Financial summary
    _heading_numbered(doc, 1, "Financial summary or highlights/Performance of the Company: (Standalone)")
    _justify(doc, f"The Company\u2019s Financial performance for the year ended on {agm['fy_end_date']} along with previous year\u2019s figures is given hereunder:")

    unit = toggles.get("FinancialUnit", "Thousand")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, f"(In {unit})", italic=True)

    fin_rows = [
        # Financial table preserves the minus sign on losses.
        ["Total Revenue", fmt_num(fin["total_revenue"]["current"], signed=True), fmt_num(fin["total_revenue"]["previous"], signed=True)],
        ["Total Expenses", fmt_num(fin["total_expenses"]["current"], signed=True), fmt_num(fin["total_expenses"]["previous"], signed=True)],
        ["Profit/Loss before Tax", fmt_num(fin["pbt"]["current"], signed=True), fmt_num(fin["pbt"]["previous"], signed=True)],
        ["Current Tax", fmt_num(fin["current_tax"]["current"], signed=True), fmt_num(fin["current_tax"]["previous"], signed=True)],
        ["Deferred Tax", fmt_num(fin["deferred_tax"]["current"], signed=True), fmt_num(fin["deferred_tax"]["previous"], signed=True)],
        ["Excess/short provision relating to earlier tax", fmt_num(fin["earlier_tax"]["current"], signed=True), fmt_num(fin["earlier_tax"]["previous"], signed=True)],
        ["Profit/Loss after Tax", fmt_num(fin["pat"]["current"], signed=True), fmt_num(fin["pat"]["previous"], signed=True)],
    ]
    add_bordered_table(doc,
        ["Particulars", f"YEAR ENDED {agm['curr_fy_end']}", f"YEAR ENDED {agm['prev_fy_end']}"],
        fin_rows, col_widths_inches=[3.2, 1.8, 1.8])

    # 2. Web Address
    _heading_numbered(doc, 2, "Web Address of Annual Return, if any:")
    _justify(doc, "Pursuant to Section 92(3) read with Rule 12 of Companies (Management and Administration) Rules, 2014 of Companies Act, 2013, (including any statutory or re-enactment thereof for the time being in force) every Company shall place a copy of the annual return on the website of the company, if any, and the web-address of such annual return shall be disclosed in the Board's report.")
    if yes(toggles, "CompanyHasWebsite") and company["website"]:
        _justify(doc, f"The web-link of the Company is {company['website']}.")
    else:
        _justify(doc, "Since, the Company doesn\u2019t have any website therefore this clause is not applicable to the Company.")

    # 3. Board Meetings
    _heading_numbered(doc, 3, "Disclosure with regard to Meetings of the Board:")
    p = doc.add_paragraph()
    add_run(p, "Board Meetings", bold=True)
    bm_count = len(data["board_meetings"])
    p = doc.add_paragraph()
    add_run(p, f"Number of meetings held: {bm_count:02d} ({_num_in_words(bm_count)})")
    _justify(doc, "The Number of meetings of the board during the year are as follows: -")
    bm_rows = [[bm["sno"], bm["date"]] for bm in data["board_meetings"]]
    add_bordered_table(doc, ["S. No", "Date of Board Meeting"], bm_rows,
                       col_widths_inches=[1.2, 2.5])

    # 4. DRS
    _heading_numbered(doc, 4, "Directors\u2019 Responsibility Statement:")
    _justify(doc, "The Directors\u2019 Responsibility Statement referred to in clause (c) of sub-section (3) of Section 134 of the Companies Act, 2013, shall state that\u2014")
    drs_points = [
        "In the preparation of the annual accounts, the applicable accounting standards had been followed along with proper explanation relating to material departures;",
        "The directors had selected such accounting policies and applied them consistently and made judgments and estimates that are reasonable and prudent so as to give a true and fair view of the state of affairs of the company at the end of the financial year and of the profit and loss of the company for that period;",
        "The directors had taken proper and sufficient care for the maintenance of adequate accounting records in accordance with the provisions of this Act for safeguarding the assets of the company and for preventing and detecting fraud and other irregularities;",
        "The directors had prepared the annual accounts on a going concern basis; and",
        "The directors had devised proper systems to ensure compliance with the provisions of all applicable laws and that such systems were adequate and operating effectively.",
    ]
    for pt in drs_points:
        p = doc.add_paragraph(pt, style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 5. Frauds
    _heading_numbered(doc, 5, "Reporting of Frauds by Auditors:")
    if yes(toggles, "FraudsReported"):
        _justify(doc, "During the Period under review the details of Instances of the fraud as reported by the Auditor under Section 143(12) of Companies Act are as follows: (Specify the Frauds reported by the auditor)")
    else:
        _justify(doc, "During the period under review, no frauds were reported by the auditors of the company under section 143(12) of the Companies Act, 2013.")

    # 6. Independent Director
    _heading_numbered(doc, 6, "Declaration by an Independent Director(s) under Section 149(6) of the Companies Act, 2013:")
    _justify(doc, "Being a Private Limited company, the company is not required to appoint any Independent Director in the Board as per the provisions of Section 149(4) of the Companies Act 2013.")

    # 7. NRC
    _heading_numbered(doc, 7, "Nomination and Remuneration Committee:")
    _justify(doc, "The provisions of Section 178(1) relating to constitution of Nomination and Remuneration Committee are not applicable to the Company and hence the Company has not devised any policy relating to appointment of Directors, payment of Managerial remuneration, Directors qualifications, positive attributes, independence of Directors and other related matters as provided under Section 178(3) of the Companies Act, 2013.")

    # 8. Auditor's Remarks
    _heading_numbered(doc, 8, "Auditor\u2019s Remarks:")
    if yes(toggles, "AuditorRemarks"):
        _justify(doc, "The Board has duly examined the Statutory Auditors\u2019 Report to the financial statements, and accordingly have provided the below mentioned comment on every observation, qualifications, or adverse remarks or disclaimer made by the auditor:")
        ar = data.get("auditor_remarks", [])
        ar_rows = (
            [[a["sno"], a["remark"], a["directors_comment"]] for a in ar]
            if ar else [["", "", ""], ["", "", ""]]
        )
        add_bordered_table(doc,
            ["S. No.",
             "Auditors\u2019 qualifications, reservations or adverse remarks or disclaimer in the auditors\u2019 report",
             "Directors\u2019 comments on qualifications, reservations or adverse remarks or disclaimer of the auditors as per Board\u2019s report"],
            ar_rows,
            col_widths_inches=[0.6, 3.2, 3.0])
    else:
        _justify(doc, "The Board has duly examined the Statutory Auditors\u2019 Report to the financial statements, which are self-explanatory and since there are no observations, qualifications, or adverse remarks or disclaimer made by the Auditors in their report, it does not call for any further explanations and comments.")

    # 9. S186
    _heading_numbered(doc, 9, "Particulars of loan, guarantee, investment or security as per Section 186:")
    is_186 = yes(toggles, "Loan186Applicable")
    s186_rows = [
        ["1.", "Whether any loan, guarantee is given by the company or securities of any other body corporate purchased?", "Yes" if is_186 else "No"],
        ["2.", "Whether the Company falls in the category provided under section 186(11)?", "Yes" if is_186 else "No"],
        ["3.", "Are there any reportable transactions on which section 186 applies? (whether or not threshold exceeds 60% of its paid-up share capital, free reserves and securities premium account or 100% of its free reserves and securities premium account)", "Yes" if is_186 else "No"],
        ["4.", "Brief details as to why transaction is not reportable", "NA"],
    ]
    add_bordered_table(doc, ["S. No.", "Particulars", "Details"], s186_rows,
                       col_widths_inches=[0.7, 4.5, 1.5])

    # 10. State of affairs \u2014 branch by profit vs loss
    _heading_numbered(doc, 10, "Brief description of state of Company\u2019s affairs:")
    pat_value = fin["pat"]["current"]
    pat_str = fmt_num(pat_value)
    is_profit = not (isinstance(pat_value, (int, float)) and pat_value < 0)
    if is_profit:
        outcome = f"earned the net profit of Rs. {pat_str}/- (In {unit})"
        _justify(doc, f"Company is engaged in the business of the {company['business_activity']} and has {outcome}.")
        _justify(doc, "That the company is exploring the new opportunities in the market for its business and your directors are taking all the necessary steps for accelerating the growth of the company. The Company will endeavor to strive to reach new heights and will toil towards attaining high profits.")
    else:
        outcome = f"incurred the net loss of Rs. {pat_str}/- (In {unit})"
        _justify(doc, f"Company is engaged in the business of the {company['business_activity']} and has {outcome}.")
        _justify(doc, "As of the reporting period, company is not operational and has not been engaged in business activities. This is a challenging time for our company, and we understand the concerns and interests of our shareholders.")

    # 11. Reserves
    _heading_numbered(doc, 11, "Transfer to reserves:")
    reserves = fin.get("reserves", "Nil")
    if str(reserves).strip().lower() in ("nil", "0", "0.00", ""):
        reserves_desc = "The Board of Directors has not proposed to transfer any amount to Reserves of the Company during the year under review"
        amt = "Nil"
    else:
        reserves_desc = f"The Board of Directors has proposed to transfer Rs. {reserves} to the Reserves of the Company during the year under review"
        amt = f"Rs. {reserves}"
    add_keyvalue_table(doc, [
        ["Brief Description", reserves_desc],
        ["Amount (in INR)", amt],
    ], label_width_inches=2.0, value_width_inches=4.5)

    # 12. Dividend
    _heading_numbered(doc, 12, "Dividend:")
    div = fin.get("dividend", "Nil")
    div_is_nil = str(div).strip().lower() in ("nil", "0", "0.00", "")
    if div_is_nil:
        if isinstance(pat_value, (int, float)) and pat_value >= 0:
            div_desc = f"There is net profit of Rs. {pat_str}/- (In {unit}) after providing for taxation during the year. The Company is in need of the liquid funds, hence, your Directors do not propose any dividend for the period under review."
        else:
            div_desc = "The Directors do not propose any dividend for the period under review."
        amt_label = "Nil"
    else:
        div_desc = f"Your Directors are pleased to recommend a dividend of Rs. {div} for the period under review."
        amt_label = f"Rs. {div}"
    add_keyvalue_table(doc, [
        ["Brief Description", div_desc],
        ["Amount (in INR)", amt_label],
    ], label_width_inches=2.0, value_width_inches=4.5)

    # 13. Material changes
    _heading_numbered(doc, 13, "Material changes and commitments occurred during the period between the end of Financial Year and the date of report, affecting financial position of company:")
    if yes(toggles, "MaterialChangesPostFY"):
        mc = data.get("material_changes", [])
        if mc:
            _justify(doc, "The following material changes and commitments have occurred during the period between the end of the financial year to which the financial statements relate and the date of this Report:")
            mc_rows = [[m["sno"], m["event_date"], m["description"]] for m in mc]
            add_bordered_table(doc,
                ["S. No.", "Date of Event", "Description of Material Change / Commitment"],
                mc_rows,
                col_widths_inches=[0.6, 1.1, 5.1])
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(4)
            add_run(p, "The following material changes and commitments have occurred during the period between the end of the financial year to which the financial statements relate and the date of this Report ")
            add_run(p, "(Specify the Changes if any)", bold=True)
    else:
        _justify(doc, "There have been no material changes and commitments affecting the financial position of the Company which have occurred during the period between the end of the financial year of the Company to which the financial statements relate and the date of the report.")

    # 14. Risk
    _heading_numbered(doc, 14, "Risk management policy:")
    if yes(toggles, "RiskMgmtPolicyExists"):
        _justify(doc, "The Company has developed and implemented a risk management policy which identifies major risks which may threaten the existence of the Company. The same has also been adopted by your Board and is also subject to its review from time to time.")
    else:
        _justify(doc, "There are no such elements of risk, which in the opinion of the Board may threaten the existence of the company during the year.")

    # 15. CSR
    _heading_numbered(doc, 15, "Corporate Social Responsibility (CSR):")
    if yes(toggles, "CSRApplicable"):
        _justify(doc, "The Company has constituted a CSR Committee in compliance with Section 135 of the Companies Act, 2013 and details of CSR activities are annexed.")
    else:
        _justify(doc, "(Not Applicable)")

    # 16. Conservation of energy
    _heading_numbered(doc, 16, "Conservation of energy, technology absorption and foreign exchange earnings and outgo as per Rule 8/8A of Companies Accounts Rules 2014:")
    add_bordered_table(doc, ["S. No.", "Particulars", "Details"],
        [["(a)", "Details regarding technology absorption as per Rule 8(3)(B)", "NA"],
         ["(b)", "Details regarding energy conservation as per Rule 8(3)(A)", "NA"],
         ["(c)", "Details regarding foreign exchange earnings and outgo as per Rule 8(3)(C)", "NA"]],
        col_widths_inches=[0.7, 4.0, 1.8])

    # 17. Subsidiaries
    _heading_numbered(doc, 17, "Performance of subsidiaries, associates and joint venture companies:")
    if yes(toggles, "HasSubsidiariesEtc"):
        _justify(doc, "The Company has the following subsidiary/Joint Venture/Associate Companies. The details are as follows:")
        add_bordered_table(doc, ["S. NO.", "Name of Company", "Subsidiary/Joint Venture/Associate Company"],
                           [["", "", ""]], col_widths_inches=[0.8, 3.0, 2.7])
    else:
        _justify(doc, "The Company does not have any subsidiary, Joint Venture or Associate Company.")

    # 18. Rule 8(5)
    _heading_numbered(doc, 18, "Disclosure as per rule 8(5) of Companies Accounts Rules 2014:")
    if yes(toggles, "SubsidiaryStatusChange"):
        rule_i = "(Insert the name of company become or ceased to be its Subsidiaries, Joint Venture or Associate Companies during the year.)"
    else:
        rule_i = "No company become or/ceased to be its Subsidiaries, Joint Venture or Associate Companies during the year."
    rule_iii = "The directors state that proper design, implementation and maintenance of adequate internal financial controls is ensured by the Company for the orderly and efficient conduct of its business, including adherence to company\u2019s policies, the safeguarding of its assets, the prevention and detection of frauds and errors, the accuracy and completeness of the accounting records, and the timely preparation of reliable financial information, as required under the Act."
    if yes(toggles, "CostRecordsApplicable"):
        rule_iv = "That disclosure regarding the maintenance of cost records as specified by the Central Government under sub-section (1) of Section 148 of Companies Act, 2013 is applicable on the Company and the company has maintained proper records and accounts of the same as required under the Act."
    else:
        rule_iv = "Pursuant to section 148 of companies Act 2013 read with Companies (Cost Records and Audit) Amendment Rules, 2014, disclosure regarding the maintenance of cost records is not applicable on the company during the period under the review."
    rule_v = ("(Specify the details and status of the application/proceeding under the Insolvency and Bankruptcy Code, 2016)"
              if yes(toggles, "InsolvencyProceedings")
              else "There is no application made or any proceeding pending under the Insolvency and Bankruptcy Code, 2016 against the company during the year.")
    rule_vi = ("(Specify the details and reasons for difference in valuation at one-time settlement vs. loan-time valuation)"
               if yes(toggles, "OneTimeSettlement")
               else "During the year under review, there has been no one-time settlement of Loan taken from banks and Financial Institutions.")
    rule85_rows = [
        ["(i)", "Disclosure of companies which have become or ceased to be its subsidiaries, joint ventures or associate companies during year", rule_i],
        ["(ii)", "Statement regarding opinion of the Board with regard to integrity, expertise and experience (including the proficiency) of the independent directors appointed during the year",
         "Being a Private Limited company, the Company is not required to appoint any Independent Director in the Board. Therefore, the Board\u2019s opinion with regard to integrity, expertise and experience (including the proficiency) of the independent directors is not required."],
        ["(iii)", "The details in respect of adequacy of internal financial controls with reference to the Financial Statements.", rule_iii],
        ["(iv)", "A disclosure, as to whether maintenance of cost records as specified by the Central Government under sub-section (1) of section 148 of the Companies Act, 2013, is required by the Company and accordingly such accounts and records are made and maintained.", rule_iv],
        ["(v)", "The details of application made or any proceeding pending under the Insolvency and Bankruptcy Code, 2016 (31 of 2016) during the year along with their status as at the end of the financial year", rule_v],
        ["(vi)", "The details of difference between amount of the valuation done at the time of one time settlement and the valuation done while taking loan from the Banks or Financial Institutions along with the reasons thereof", rule_vi],
    ]
    add_bordered_table(doc, ["S. No.", "Particulars", "Details"], rule85_rows,
                       col_widths_inches=[0.6, 2.7, 3.2])

    # 19. Highlights — branch label by profit/loss
    _heading_numbered(doc, 19, "Financial summary or highlights:")
    pat_cur = fin["pat"]["current"]
    is_loss = isinstance(pat_cur, (int, float)) and pat_cur < 0
    if is_loss:
        _justify(doc,
            f"During the Financial Year {agm['fy_label']}, revenue from operations of the company is Rs. "
            f"{fmt_num(fin['total_revenue']['current'])}/- (In {unit}) and Net Loss is Rs. "
            f"{fmt_num(pat_cur)}/- (In {unit}).")
    else:
        _justify(doc,
            f"During the Financial Year {agm['fy_label']}, revenue from operations of the company is Rs. "
            f"{fmt_num(fin['total_revenue']['current'])}/- (In {unit}) and Net Profit after Tax (PAT) is Rs. "
            f"{fmt_num(pat_cur)}/- (In {unit}).")

    # 20. Change in nature
    _heading_numbered(doc, 20, "Change in the nature of business, if any:")
    cibn_yes = yes(toggles, "ChangeInBusinessNature")
    bnc_rows_in = data.get("business_nature_changes", []) if cibn_yes else []
    # If toggle is YES but no descriptions are filled, leave a placeholder.
    if cibn_yes:
        if bnc_rows_in:
            description_cell = "; ".join(r["description"] for r in bnc_rows_in)
        else:
            description_cell = "(Insert the description of the nature of business changed)"
    else:
        description_cell = "NA"
    cibn_rows = [
        ["1.", "Whether there has been any change in the nature of business during the year",
         "Yes" if cibn_yes else "No"],
        ["2.", "Brief description of the change (if applicable)", description_cell],
    ]
    add_bordered_table(doc, ["S. No.", "Particulars", "Details"], cibn_rows,
                       col_widths_inches=[0.6, 4.0, 2.2])

    # 21. Directors
    _heading_numbered(doc, 21, "Directors and Key Managerial Personnel:")
    if yes(toggles, "ChangeInBoardComposition"):
        _justify(doc, "There is change in the composition of the board of directors of the company during the year.")
    else:
        _justify(doc, "There is no change in the composition of board of directors during the year.")
    _justify(doc, "The composition of the Board of Directors are as follows:")
    dir_rows = [[d["sno"], d["din"], d["name"], d["appointment"], d["cessation"] or "-"] for d in data["directors"]]
    add_bordered_table(doc, ["S. No.", "DIN", "Name of Director", "Date of appointment", "Date of cessation (if any)"],
                       dir_rows, col_widths_inches=[0.6, 1.0, 2.4, 1.3, 1.2])

    # Branch the regularization sentence by appointment date vs FY end.
    from datetime import datetime as _dt
    def _parse_ddmmyyyy(s):
        try:
            return _dt.strptime(str(s).strip(), "%d/%m/%Y")
        except (ValueError, TypeError):
            return None
    fy_end_dt = _parse_ddmmyyyy(agm.get("curr_fy_end", ""))

    def _date_short_words(s):
        d = _parse_ddmmyyyy(s)
        if d is None: return s or ""
        # ordinal suffix
        n = d.day
        if 10 <= n % 100 <= 20:
            suf = "th"
        else:
            suf = {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
        return f"{n}{suf} {d.strftime('%B')}, {d.year}"

    for dc in regs:
        add_para(doc, "")
        appt_dt = _parse_ddmmyyyy(dc.get("appt_date", ""))
        appt_words = _date_short_words(dc.get("appt_date", ""))
        agm_words = agm.get("date_words") or _date_short_words(agm.get("date", ""))
        # If appointment is AFTER FY end → "After the end of financial year, …"
        # else (during the FY) → "During the year, …"
        if fy_end_dt and appt_dt and appt_dt > fy_end_dt:
            lead = "After the end of financial year"
        else:
            lead = "During the year"
        _justify(doc,
            f"{lead}, {dc['name']} (DIN: {dc['din']}) has been appointed as an Additional "
            f"Director of the Company in the Board Meeting held on {appt_words} who will be "
            f"regularised in the upcoming Annual General Meeting to be held on {agm_words}.")

    # 22. Deposits
    _heading_numbered(doc, 22, "Deposits:")
    _justify(doc, "No deposit has been accepted by the company during the year. Therefore, the disclosures specified under Chapter V of Rule 8 (5) of Companies (Accounts) Rules, 2014 is not applicable on the company.")

    # 23. Significant orders
    _heading_numbered(doc, 23, "Details of significant and material orders passed by the regulators or courts or tribunals impacting the going concern status and company\u2019s operations in future:")
    if yes(toggles, "SignificantOrdersByRegulators"):
        _justify(doc, "The details of a significant material order passed by the regulators/Courts/Tribunals which may impact the going concern status of the Company and its future operations is provided as under: (Specify the Details of orders)")
    else:
        _justify(doc, "No significant and material order has been passed by the regulators, courts, tribunals impacting the going concern status and Company\u2019s operations in future")

    # 24. Annual evaluation
    _heading_numbered(doc, 24, "Annual Evaluation:")
    _justify(doc, "Being a Private Limited company, the provisions of section 134(3)(p) relating to Annual Evaluation of Performance of the board its committees and of the individual directors is not applicable on the company.")

    # 25. POSH
    _heading_numbered(doc, 25, "Obligation of Company under the Sexual Harassment of Women At Workplace (Prevention, Prohibition And Redressal) Act, 2013:")
    _justify(doc, "The Company recognizes its duty to provide safe and secure working environment at the workplace and thus, in line with the requirements of the Sexual Harassment of Women at the workplace (Prevention, Prohibition and Redressal) Act, 2013, the Company has in place a policy for prevention of Sexual Harassment of women at the workplace and has also set up an Internal Complaints Committee (ICC) to redress complaints received regarding sexual harassment. All employees are covered under this policy.")
    _justify(doc, f"Your Board takes pride in presenting the summary of sexual harassment complaints received and disposed of during the year {agm['fy_label']}:")
    add_bordered_table(doc, ["S. No.", "Particulars", "Details"],
        [["(i)", "Number of Sexual Harassment Complaints received", str(emp["sh_received"])],
         ["(ii)", "Number of Sexual Harassment Complaints disposed off", str(emp["sh_disposed"])],
         ["(iii)", "Number of Sexual Harassment Complaints pending beyond 90 days.", str(emp["sh_pending"])]],
        col_widths_inches=[0.7, 4.0, 1.8])

    # 26. Maternity — auto-derived from female employee count.
    # An explicit MaternityActApplicable=YES toggle can still force the
    # "applicable" branch (e.g. for companies expecting female hires), but
    # if there's at least one female employee the Act is applicable by law,
    # so we never claim "no female employee" while the count is > 0.
    _heading_numbered(doc, 26, "Disclosure as per Maternity Benefit Act, 1961:")
    has_female_emp = int(emp.get("female", 0) or 0) > 0
    if has_female_emp or yes(toggles, "MaternityActApplicable"):
        _justify(doc, "The Company complies with the provisions of the Maternity Benefit Act, 1961. The Company ensures that all female employees are provided with maternity leaves, benefits and protections as mandated by the Act.")
    else:
        _justify(doc, "As the Company does not have any female employee. Therefore, the provisions of the Maternity Benefit Act, 1961 are not applicable to the Company.")

    # 27. Employees — Transgender row is suppressed when the count is 0 / blank.
    # Most small private companies don't have any transgender employees, and the
    # user asked for the row to be hidden in that case to avoid an awkward empty
    # statutory line.
    _heading_numbered(doc, 27, "Number of Employees:")
    emp_rows = [
        ["1.", "Female", str(emp["female"])],
        ["2.", "Male", str(emp["male"])],
    ]
    if int(emp.get("transgender", 0) or 0) > 0:
        emp_rows.append([str(len(emp_rows) + 1) + ".", "Transgender", str(emp["transgender"])])
    add_bordered_table(doc, ["S. No.", "Category of Employees", "Number of Employees"],
        emp_rows, col_widths_inches=[0.7, 3.0, 2.5])

    # 28. Others
    _heading_numbered(doc, 28, "Others:")

    _heading_lettered(doc, "1", "Change of Name:")
    if yes(toggles, "ChangeOfNameDuringYear"):
        old = company.get("old_name") or "(Existing Name)"
        # Date resolution order:
        #   1) Single 'EGM Date' field on the Company tab
        #   2) First EGMMeetings row
        #   3) First MaterialChanges row whose description mentions "name"
        #   4) Placeholder underscores
        egm_date = agm.get("egm_date")
        if not egm_date:
            egms = data.get("egm_meetings", [])
            if egms:
                egm_date = egms[0]["date"]
        if not egm_date:
            for mc in data.get("material_changes", []):
                if "name" in (mc.get("description") or "").lower():
                    egm_date = mc.get("event_date")
                    break
        egm_date = egm_date or "____________"
        _justify(doc, f"During the year company has changed its name from {old} to {company['name']} in the duly convened Extra-Ordinary General Meeting held on {egm_date}.")
    else:
        _justify(doc, "The Company has not changed its name within the financial year.")

    _heading_lettered(doc, "2", "Share Capital:")
    scc = data.get("share_capital_changes", []) if yes(toggles, "ChangeInShareCapital") else []
    if yes(toggles, "ChangeInShareCapital"):
        _justify(doc, "During the year under review, there has been a change in the Share Capital of the Company.")
    else:
        _justify(doc, "During the year under review, there has been no change in the Share Capital of the Company.")
    _justify(doc, f"The Share Capital of the Company as on {agm['fy_end_date']}:")
    sc_rows = [
        ["a) Authorized Capital:", sc["authorized"]],
        ["b) Issued Capital:", sc["issued"]],
        ["c) Subscribed and Paid-up Capital:", sc["paid_up"]],
    ]
    add_bordered_table(doc, ["Share capital", "Description Of Capital"], sc_rows,
        col_widths_inches=[2.5, 4.0])
    for s in scc:
        _justify(doc, s["description"])
    additional_capital = (sc.get("additional") or "").strip()
    if additional_capital:
        add_para(doc, "")
        _justify(doc, additional_capital)
    _justify(doc, "During the year, the Company has not issued any equity share with differential voting rights hence the disclosure under Rule 4(4) of the Companies (Share Capital and Debentures) Rules, 2014 is not applicable.")

    _heading_lettered(doc, "3", "General Meetings:")
    p = doc.add_paragraph()
    add_run(p, "a) Annual General Meeting:", bold=True, underline=True)
    # A non-empty Previous AGM Date is the signal that a previous AGM was held.
    # 1st AGM = leave Previous AGM Date blank.
    prev_agm_date = agm.get("prev_agm_date", "").strip()
    if prev_agm_date:
        _justify(doc, f"During the period under review the Company has held its Annual General Meeting on {prev_agm_date} in compliance with the provisions of Companies Act, 2013 and SS-2 (Secretarial Standard on General Meetings) issued by The Institute of Company Secretaries of India (ICSI).")
    else:
        _justify(doc, f"During the period under review the Company has not held any Annual General Meeting as the Company was incorporated on {company['incorporation_date']}.")
    p = doc.add_paragraph()
    add_run(p, "b) Extra Ordinary General Meeting:", bold=True, underline=True)
    if yes(toggles, "HeldEGM"):
        _justify(doc,
            "During the period under review the company has held the Extra Ordinary General Meeting "
            "in compliance with the provisions of the Companies Act 2013, SS-2 (Secretarial Standard "
            "on General Meetings) issued by The Institute of Company Secretaries of India (ICSI) "
            "(including any statutory or re-enactment thereof for the time being in force).")
        egm_rows_in = data.get("egm_meetings", [])
        # Only render the table if the user actually filled rows on the
        # EGMMeetings tab. Empty → suppress the table entirely.
        if egm_rows_in:
            egm_rows = [[r["sno"] or i + 1, r["date"], r["description"] or "-"]
                        for i, r in enumerate(egm_rows_in)]
            add_bordered_table(doc, ["S. No.", "Date of Meeting", "Description"], egm_rows,
                               col_widths_inches=[0.55, 1.5, 4.55])
    else:
        _justify(doc, "During the period under review the company had not held any Extra Ordinary General Meeting.")

    _heading_lettered(doc, "4", "Auditors:")
    p = doc.add_paragraph()
    add_run(p, "Statutory Auditor:-", bold=True, underline=True)
    # Per reference template, the "Pursuant to the provisions of Section 139..."
    # paragraph appears whenever the auditor block is shown — either because the
    # existing auditor is being mentioned (AuditorReportInDirReport) OR because
    # a re-appointment is being proposed at this AGM (AuditorReappointment).
    if yes(toggles, "AuditorReportInDirReport") or yes(toggles, "AuditorReappointment"):
        first_term_end_year = _fy_end_year(agm.get("curr_fy_end", ""))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)
        add_run(p, "Pursuant to the provisions of Section 139 of Companies Act, 2013 read with the Companies (Audit and Auditors) Rules, 2014, (including any statutory or re-enactment thereof for the time being in force) ")
        add_run(p, f"{aud['firm_name']}, {aud['designation']} (Firm Registration No. {aud['frn']})", bold=True)
        add_run(p, f", the Statutory Auditors of the Company, shall hold office up to the conclusion of Annual General Meeting held in F.Y. {first_term_end_year}.")
    if yes(toggles, "AuditorReappointment"):
        try:
            term = int(str(aud.get("term_years", "5")).strip() or "5")
        except (ValueError, TypeError):
            term = 5
        try:
            start_year = int(_fy_end_year(agm.get("curr_fy_end", "")))
            term_phrase = (f"{term:02d} Financial Years i.e. "
                           f"{start_year}-{start_year + 1} to "
                           f"{start_year + term - 1}-{start_year + term}")
        except (ValueError, TypeError):
            term_phrase = f"{term:02d} Financial Years"
        agm_date_long = _format_date_long(agm.get("date", ""))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)
        add_run(p, f"In the upcoming Annual General Meeting going to be held on {agm['day'].title()}, {agm_date_long}, ")
        add_run(p, f"{aud['firm_name']}, {aud['designation']} (Firm Registration No. {aud['frn']})", bold=True)
        add_run(p, f" will be proposed to be re-appointed as the Statutory Auditors of the Company for the {term_phrase}.")
        _justify(doc, "The Company has received a certificate from the said Auditors that they are eligible to hold office as the Auditors of the Company and are not disqualified for being so appointed.")

    p = doc.add_paragraph()
    add_run(p, "Secretarial Auditor:-", bold=True, underline=True)
    if yes(toggles, "SecretarialAuditorRequired"):
        _justify(doc, "The Company has appointed a Secretarial Auditor in compliance with Section 204 of the Companies Act, 2013.")
    else:
        _justify(doc, "That Pursuant to section 204 of companies Act 2013 and rules made thereunder the requirement to appoint the Secretarial auditor is not applicable on the company during the year.")

    p = doc.add_paragraph()
    add_run(p, "Cost Auditor:-", bold=True, underline=True)
    if yes(toggles, "CostAuditorRequired"):
        _justify(doc, "The Company has appointed a Cost Auditor in compliance with Section 148 of the Companies Act, 2013.")
    else:
        _justify(doc, "That Pursuant to section 148 of companies Act 2013 and rules made thereunder the requirement to appoint the Cost Auditor is not applicable on the company during the year.")

    _heading_lettered(doc, "5", "Vigil Mechanism:")
    if yes(toggles, "VigilMechanismRequired"):
        _justify(doc, "The Company has established a vigil mechanism for directors and employees in compliance with Section 177(10) of the Companies Act, 2013.")
    else:
        _justify(doc, "That the Company is not required to establish the vigil mechanism for directors and employees to report the genuine concerns hence the disclosures under section 177(10) of the Companies Act, 2013 are not applicable on the company.")

    _heading_lettered(doc, "6", "Particulars of contracts or arrangements with related parties:")
    if yes(toggles, "RPT_Required"):
        _justify(doc, "During the year under review, the Company entered into contracts or arrangements with its related parties referred to in Section 188(1) of the Companies Act, 2013. Disclosures in Form AOC-2 in terms of Section 134(3)(h) of the Companies Act, 2013 and Rule 8 of the Companies (Accounts) Rules, 2014 is included in this report as (\u201cAnnexure-A\u201d) and forms an integral part of this report.")
    else:
        _justify(doc, "During the year under review, the Company did not enter into any contracts or arrangements with its related parties referred to in Section 188(1) of the Companies Act, 2013.")

    _heading_lettered(doc, "7", "Corporate Governance Certificate:")
    _justify(doc, "(Not Applicable)")

    _heading_lettered(doc, "8", "Management Discussion and Analysis:")
    _justify(doc, "(Not Applicable)")

    _heading_lettered(doc, "9", "Human Resources:")
    _justify(doc, "Company treats its \u201chuman resources\u201d as one of its most important assets. Company continuously invests in attraction, retention and development of talent on an ongoing basis. A number of programs that provide focused people attention are currently underway. Company thrust is on the promotion of talent internally through job rotation and job enlargement.")

    _heading_lettered(doc, "10", "Transfer of Amounts to Investor Education and Protection Fund:")
    _justify(doc, "Company did not have any funds lying unpaid or unclaimed for a period of seven years. Therefore, there were no funds which were required to be transferred to Investor Education and Protection Fund (IEPF).")

    _heading_lettered(doc, "11", "Secretarial Standards:")
    _justify(doc, "The Company has compiled with Secretarial Standards Issued by the Institute of Company Secretaries of India on Board and General Meetings.")

    _heading_lettered(doc, "12", "Appointment of Designated Person as per Rule 9 of Companies (Management and Administration) Rules, 2014:")
    _justify(doc, "In accordance with Rule 9 of Companies (Management and Administration) Rules, 2014, it is essential for the Company to designate a person who shall be responsible for furnishing, and extending co-operation for providing, information to the Registrar or any other authorised officer with respect to beneficial interest in shares of the company.")
    _justify(doc, f"The Company has appointed {signing['designated_person']['name']} (DIN: {signing['designated_person']['din']}), Director as Designated Person in the Board Meeting and the same has been reported in the Annual Return of the Company.")

    _heading_lettered(doc, "13", "Acknowledgements:")
    _justify(doc, "Your directors place on the record their appreciation of the contribution made by employees, consultants at all levels, who with their competence, diligence, solidarity, co-operation and support have enabled the Company to achieve the desired results.")
    _justify(doc, "The board of Directors gratefully acknowledge the assistance and co-operation received from the Central and State Governments Departments, Shareholders and Stakeholders.")

    add_two_signer_block(doc, signing["second"], signing["first"],
                         date_str=agm["notice_date"], place_str=agm["notice_place"],
                         company_name=company["name"])


# =====================================================================
# SECTION 3: AOC-2
# =====================================================================
def section_aoc2(doc, data):
    add_centered_heading(doc, "FORM NO. AOC-2", size=12, underline=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "(Pursuant to clause (h) of sub-section (3) of section 134 of the Act and Rule 8(2) of the Companies (Accounts) Rules, 2014.)", italic=True)

    add_para(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "Form for Disclosure of particulars of contracts/arrangements entered into by the company with related parties referred to in sub-section (1) of section 188 of the Companies Act, 2013 including certain arm\u2019s length transactions under third proviso thereto.")

    add_para(doc, "")
    p = doc.add_paragraph()
    add_run(p, "*Name of the Company: ", bold=True)
    add_run(p, data["company"]["name"], bold=True)

    rpts = data.get("rpts", []) if yes(data["toggles"], "RPT_Required") else []

    # ---- Section 1: Not at arm's length ----
    add_para(doc, "")
    p = doc.add_paragraph()
    add_run(p, "1. Details of contracts or arrangements or transactions not at Arm\u2019s length basis.", bold=True)

    add_aoc2_count_box(
        doc,
        "*Number of contracts or arrangements or transactions not at arm\u2019s length basis",
        "0",
    )
    add_para(doc, "")

    not_arms_particulars = [
        ("a)", "Corporate identity number (CIN) or foreign company registration number (FCRN) or Limited Liability Partnership number (LLPIN) or Foreign Limited Liability Partnership number (FLLPIN) or Permanent Account Number (PAN)/Passport for individuals or any other registration number"),
        ("b)", "Name (s) of the related party"),
        ("c)", "Nature of relationship"),
        ("d)", "Nature of contracts/arrangements/transaction"),
        ("e)", "Duration of the contracts/arrangements/transaction"),
        ("f)", "Salient terms of the contracts or arrangements or transaction including the value, if any"),
        ("g)", "Justification for entering into such contracts or arrangements or transactions"),
        ("h)", "Date of approval by the Board"),
        ("i)", "Amount paid as advances, if any"),
        ("j)", "Date on which the special resolution was passed in General meeting as required under first proviso to section 188"),
        ("k)", "SRN of MGT-14"),
    ]
    add_aoc2_block_table(doc, "Block-1",
        [(letter, desc, "") for letter, desc in not_arms_particulars],
        col_widths_inches=[0.5, 4.2, 1.8])

    # ---- Section 2: At arm's length ----
    add_para(doc, "")
    p = doc.add_paragraph()
    add_run(p, "2. Details of contracts or arrangements or transactions at Arm\u2019s length basis.", bold=True)

    add_aoc2_count_box(
        doc,
        "Number of material contracts or arrangements or transactions at arm\u2019s length basis",
        str(len(rpts)),
    )

    for idx, rpt in enumerate(rpts, start=1):
        add_para(doc, "")
        rpt_titled_name = _rpt_titled(rpt)
        arms_rows = [
            ("a)", "Corporate identity number (CIN) or foreign company registration number (FCRN) or Limited Liability Partnership number (LLPIN) or Foreign Limited Liability Partnership number (FLLPIN) or Permanent Account Number (PAN)/Passport for individuals or any other registration number", rpt["pan_cin"]),
            ("b)", "Name (s) of the related party", rpt_titled_name),
            ("c)", "Nature of relationship", rpt["relationship"]),
            ("d)", "Nature of contracts/arrangements/transaction", rpt["nature_of_contract"]),
            ("e)", "Duration of the contracts/arrangements/transaction", rpt["duration"]),
            ("f)", "Salient terms of the contracts or arrangements or transaction including the value, if any", rpt["salient_terms"]),
            ("g)", "Date of approval by the Board", rpt["board_approval_date"]),
            ("h)", "Amount paid as advances, if any", rpt["advances"] or "Nil"),
            ("i)", "Date on which the special resolution was passed in General meeting as required under first proviso to section 188",
                   rpt.get("special_resolution_date") or "NA"),
            ("j)", "SRN of MGT-14", rpt.get("mgt14_srn") or "NA"),
        ]
        add_aoc2_block_table(doc, f"Block-{idx}", arms_rows,
                             col_widths_inches=[0.5, 4.2, 1.8])

    add_para(doc, "")
    add_two_signer_block(doc,
        first=data["signing"]["second"], second=data["signing"]["first"],
        date_str=data["agm"]["notice_date"], place_str=data["agm"]["notice_place"],
        company_name=data["company"]["name"],
    )


# =====================================================================
# SECTION 4: SHAREHOLDERS LIST (table page)
# =====================================================================
def section_shareholders_list(doc, data):
    """Equity list of shareholders - table embedded in Word."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"LIST OF SHAREHOLDERS AS ON {data['agm']['fy_end_date'].upper()}", bold=True, size=12, underline=True)

    p = doc.add_paragraph()
    add_run(p, "TYPE OF SHARES: EQUITY                          AMOUNT PER SHARE (In Rs): 10.00/-", bold=True)
    add_para(doc, "")

    headers = [
        "Sr. No.", "Type of shareholder/ debenture holder", "Category of shareholder",
        "Details of shareholder/ debenture holder", "Name of shareholder/ debenture holder",
        "Type of security held", "Class of security held", "Folio Number / Reference Number",
        "Nationality/ Country of incorporation", "Gender", "Type of Identifier",
        "Identification No.", "Occupation", "Number of security held",
        "Nominal value per security", "Total amount of securities held (in INR)",
    ]
    rows = []
    total_shares = 0
    total_amount = 0
    for sh in data["shareholders"]:
        try:
            total_shares += int(float(sh["shares"]) or 0)
        except (ValueError, TypeError):
            pass
        try:
            total_amount += float(sh["total"]) or 0
        except (ValueError, TypeError):
            pass
        rows.append([
            sh["sno"],
            sh["type"],
            sh["category"],
            "Not applicable",
            sh["name"],
            "Equity",
            "1",
            sh["folio"],
            "India",
            sh.get("gender", ""),
            "Income Tax PAN",
            sh["pan"],
            sh["occupation"],
            sh["shares"],
            sh["nominal"],
            sh["total"],
        ])
    # Totals row — only the two numeric columns at the end populated; others blank.
    rows.append(["", "", "", "", "Total", "", "", "", "", "", "", "", "",
                 total_shares, "", f"{int(total_amount):,}"])

    tbl = add_bordered_table(doc, headers, rows,
        col_widths_inches=[0.35, 0.75, 0.65, 0.65, 1.05, 0.55, 0.45, 0.6, 0.7, 0.45, 0.6, 0.85, 0.6, 0.55, 0.6, 0.75])
    # Compact font + bold the totals row.
    for r_idx, row in enumerate(tbl.rows):
        is_totals = (r_idx == len(tbl.rows) - 1)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    if is_totals:
                        run.font.bold = True

    add_para(doc, "")
    add_two_signer_block(doc,
        first=data["signing"]["second"], second=data["signing"]["first"],
        date_str="", place_str=data["agm"]["notice_place"],
        company_name=data["company"]["name"],
    )


# =====================================================================
# SECTION 5: ATTENDANCE OF AGM (table page)
# =====================================================================
def section_attendance_agm(doc, data):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"ATTENDANCE SHEET OF MEMBERS IN AGM FOR FY {data['agm']['fy_label']}",
            bold=True, size=12, underline=True)

    add_para(doc, "")
    headers = ["S. No.", "MEMBER\u2019S NAME", "PROXY NAME", "SIGNATURE"]
    rows = [[sh["sno"], sh["name"], "", ""] for sh in data["shareholders"]]
    add_bordered_table(doc, headers, rows, col_widths_inches=[0.7, 2.3, 2.0, 2.0])


# =====================================================================
# SECTION 6: LIST OF DIRECTORS (table page)
# =====================================================================
def section_directors_list(doc, data):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"LIST OF DIRECTORS AS ON {data['agm']['fy_end_date'].upper()}",
            bold=True, size=12, underline=True)

    add_para(doc, "")
    # Filter: only directors whose appointment date is ≤ FY end. Anyone
    # appointed AFTER FY end wasn't on the board "as on" that date.
    from datetime import datetime as _dt
    def _pdt(s):
        try: return _dt.strptime(str(s).strip(), "%d/%m/%Y")
        except (ValueError, TypeError): return None
    fy_end = _pdt(data["agm"].get("curr_fy_end", ""))
    eligible = []
    for d in data["directors"]:
        appt = _pdt(d.get("appointment", ""))
        if fy_end and appt and appt > fy_end:
            continue
        eligible.append(d)

    headers = ["S. No.", "DIN", "Name of Director", "Date of Appointment", "Date of Cessation"]
    rows = [[i, d["din"], d["name"], d["appointment"], d["cessation"] or "-"]
            for i, d in enumerate(eligible, start=1)]
    tbl = add_bordered_table(doc, headers, rows,
                              col_widths_inches=[0.6, 1.2, 3.0, 1.2, 1.2])
    for row in tbl.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    add_para(doc, "")
    add_two_signer_block(doc,
        first=data["signing"]["second"], second=data["signing"]["first"],
        date_str="", place_str=data["agm"]["notice_place"],
        company_name=data["company"]["name"],
    )


# =====================================================================
# SECTION 7: DIRECTOR ATTENDANCE (table page)
# =====================================================================
def section_director_attendance(doc, data):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"Director's Attendance for the F.Y. {data['agm']['fy_label']}",
            bold=True, size=12, underline=True)

    add_para(doc, "")
    headers = ["S. No.", "Name of Director", "DIN"] + [bm["date"] for bm in data["board_meetings"]]
    rows = []
    for d_idx, d in enumerate(data["directors"]):
        attendance_cells = []
        for bm in data["board_meetings"]:
            att = bm.get("attendance", [])
            # Show exactly what the user typed: empty → empty, "N/A" → "N/A",
            # "P" → "P", etc. No silent conversion.
            val = att[d_idx].strip() if d_idx < len(att) else ""
            attendance_cells.append(val)
        row = [d["sno"], d["name"], d["din"]] + attendance_cells
        rows.append(row)

    # Shrink S.No/DIN aggressively + use the slack from the tightened landscape
    # margins (10.69" usable) to give date columns enough room for DD/MM/YYYY
    # on a single line, even for 10 board meetings.
    n_meet = len(data["board_meetings"]) or 1
    fixed_w = 0.30 + 1.25 + 0.75    # S.No + Name + DIN = 2.30"
    avail_for_dates = 10.65 - fixed_w  # ≈ 8.35"
    date_col_w = max(0.80, avail_for_dates / n_meet)
    col_widths = [0.30, 1.25, 0.75] + [date_col_w] * n_meet
    tbl = add_bordered_table(doc, headers, rows, col_widths_inches=col_widths)
    # Make the signing cells taller so directors can sign
    for r_idx in range(1, len(rows) + 1):
        tbl.rows[r_idx].height = Inches(0.5)
    # 8pt across the board — keeps every date on one line.
    for row in tbl.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    add_para(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "Note: Above declaration is given by directors of the company that they were presented in the meeting of Board of Directors on above mentioned dates.",
            bold=True, size=10)


# =====================================================================
# SECTION 8: RESOLUTION - DIRECTOR REGULARIZATION
# =====================================================================
def section_resolution_director(doc, data):
    if not yes(data["toggles"], "RegulariseAdditionalDirector"):
        return
    regs = data.get("regularizations", [])
    if not regs:
        return
    agm = data["agm"]
    signing = data["signing"]
    company = data["company"]

    for idx, dc in enumerate(regs):
        if idx > 0:
            _new_page_with_letterhead(doc, data)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, "Certified True Copy of the Resolution Passed in the ", bold=True)
        add_run(p, f"{agm['number']} Annual General Meeting of Members held on ", bold=True)
        add_run(p, f"{agm['day'].title()}, {agm['date_words']} ", bold=True)
        add_run(p, "at ", bold=True)
        add_run(p, f"{agm['venue']} ", bold=True)
        add_run(p, "at ", bold=True)
        add_run(p, f"{agm['time']}", bold=True)

        add_para(doc, "")
        p = doc.add_paragraph()
        add_run(p, f"To regularize the appointment of {dc['name']} (DIN: {dc['din']}), as the Director of the Company", bold=True)

        add_para(doc, "")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, "\u201cResolved that ", bold=True)
        add_run(p, "in accordance with the provision of Section 161(1), read with Section 152 of the Companies Act, 2013 and the Rules made there under (including any statutory modification(s) or re-enactment thereof), and the Article of Association of the Company, ")
        add_run(p, f"{dc['name']} (DIN: {dc['din']}), ", bold=True)
        add_run(p, "be and is hereby appointed as a Director of the Company who was appointed as an Additional Director of the Company by the Board of Directors at its meeting held on ")
        add_run(p, f"{dc['appt_dotted']}", bold=True)
        add_run(p, ".\u201d")

        add_para(doc, "")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, "\u201cResolved Further That ", bold=True)
        add_run(p, "any Director of the Company of the company be and are hereby individually/severally authorized to digitally sign and file e-form DIR-12 with the Registrar of Companies and to do all such things, deeds, acts which may deem necessary to give effect of the aforesaid resolution.\u201d")

        add_para(doc, "")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, "// Certified True Copy //", bold=True)

        add_two_signer_block(doc, signing["second"], signing["first"],
                             date_str=agm["date"], place_str=agm["notice_place"],
                             company_name=company["name"])


# =====================================================================
# SECTION 9: RESOLUTION - RPT
# =====================================================================
def section_resolution_rpt(doc, data):
    if not yes(data["toggles"], "RPT_Required"):
        return
    rpts = data.get("rpts", [])
    if not rpts:
        return
    agm = data["agm"]
    signing = data["signing"]
    company = data["company"]

    for idx, rpt in enumerate(rpts):
        if idx > 0:
            _new_page_with_letterhead(doc, data)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, "Certified True Copy of the Resolution Passed in the ", bold=True)
        add_run(p, f"{agm['number']} Annual General Meeting of Members held on ", bold=True)
        add_run(p, f"{agm['day'].title()}, {agm['date_words']} ", bold=True)
        add_run(p, "at ", bold=True)
        add_run(p, f"{agm['venue']} ", bold=True)
        add_run(p, "at ", bold=True)
        add_run(p, f"{agm['time']}", bold=True)

        add_para(doc, "")
        p = doc.add_paragraph()
        add_run(p, f"To approve the related party transactions under section 188 of the Companies Act, 2013 with {_rpt_titled(rpt)}.", bold=True)

        add_para(doc, "")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, "\u201cResolved that ", bold=True)
        add_run(p, "pursuant to the provisions of Section 188 of the Companies Act, 2013 (\u201cAct\u201d) and other applicable provisions, if any, read with Rule 15 of the Companies (Meetings of Board and its Powers) Rules, 2014, (including any amendments, modifications, variations or re-enactments thereof for the time being in force), the consent of the members of the Company be and is hereby accorded for carrying out and / or continuing with arrangements and transactions with ")
        add_run(p, f"{_rpt_titled(rpt)}, ", bold=True)
        add_run(p, f"related party of the Company with respect to {rpt['nature_of_contract'].lower()}, at arm\u2019s length basis and in the ordinary course of business, on such terms and conditions as the Board of Directors may deem fit up to a maximum aggregate value of ")
        add_run(p, _currency_text(rpt["max_value"]), bold=True)
        add_run(p, " each.")

        add_para(doc, "")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, "Resolved further that ", bold=True)
        add_run(p, "for the purpose of giving effect to the above resolution, the Board of Directors of the Company be and are hereby authorized to do all acts, deeds and things in their absolute discretion that may be considered necessary, proper and expedient or incidental for the purpose of giving effect to this resolution in the interest of the Company.\u201d")

        add_para(doc, "")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, "//Certified true copy//", bold=True)

        add_two_signer_block(doc, signing["second"], signing["first"],
                             date_str=agm["date"], place_str=agm["notice_place"],
                             company_name=company["name"])


# =====================================================================
# SECTION 10: RESOLUTION - AUDITOR
# =====================================================================
def section_resolution_auditor(doc, data):
    if not yes(data["toggles"], "AuditorReappointment"):
        return
    aud = data["auditor"]
    agm = data["agm"]
    signing = data["signing"]
    company = data["company"]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "Certified True Copy of the Resolution Passed in the ", bold=True)
    add_run(p, f"{agm['number']} Annual General Meeting of Members held on ", bold=True)
    add_run(p, f"{agm['day'].title()}, {agm['date_words']} ", bold=True)
    add_run(p, "at ", bold=True)
    add_run(p, f"{agm['venue']} ", bold=True)
    add_run(p, "at ", bold=True)
    add_run(p, f"{agm['time']}", bold=True)

    add_para(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "\u201cRESOLVED THAT ", bold=True)
    add_run(p, "pursuant to the provisions of Section 139 and other applicable provisions, if any, of the Companies Act, 2013 and the Rules framed there under, as amended from time to time, ")
    add_run(p, f"{aud['firm_name']}, {aud['designation']} (Firm Registration No. {aud['frn']}) ", bold=True)
    add_run(p, f"be and are hereby re-appointed as Auditors of the Company to hold office from the conclusion of this Annual General Meeting till the conclusion of the Annual General Meeting held in F.Y. {aud['tenure_end_fy']} of the Company, at such remuneration plus service tax, out of pocket, travelling and living expenses, etc., as may be mutually agreed between the Board of Directors of the Company and the Auditors.\u201d")

    add_para(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "\u201cRESOLVED FURTHER THAT ", bold=True)
    add_run(p, "any Director of the Company be and are hereby individually/severally authorized to digitally sign and file e-form ADT-1 with the Registrar of Companies and to do all such things, deeds, acts which may deem necessary to give effect of the aforesaid resolution.\u201d")

    add_para(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "//Certified True Copy//", bold=True)

    add_single_signer_block(doc, signer=signing["first"],
                             date_str=agm["date"], place_str=agm["notice_place"],
                             closing="", company_name=company["name"])


# =====================================================================
# SECTION 11: INTIMATION LETTER
# =====================================================================
def section_intimation_letter(doc, data):
    if not yes(data["toggles"], "AuditorReappointment"):
        return
    aud = data["auditor"]
    agm = data["agm"]
    signing = data["signing"]
    company = data["company"]

    # Compact letter — 10pt body and tight spacing so it fits on one page.
    BODY = 10
    add_centered_heading(doc, "INTIMATION LETTER", size=12, underline=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(3)
    add_run(p, f"Date: {agm['notice_date']}", bold=True, size=BODY)

    # Tight address block — no gap between To / firm / designation / address lines.
    def _tight(p):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        return p

    p = _tight(doc.add_paragraph()); add_run(p, "To,", bold=True, size=BODY)
    p = _tight(doc.add_paragraph()); add_run(p, f"{aud['firm_name']}", bold=True, size=BODY)
    p = _tight(doc.add_paragraph()); add_run(p, f"{aud['designation']},", bold=True, size=BODY)
    p = _tight(doc.add_paragraph())
    p.paragraph_format.space_after = Pt(6)
    add_run(p, f"ADDRESS: {aud['address']}", bold=True, size=BODY)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_run(p, "Sub: Eligibility for Re-Appointment as Statutory Auditors of the Company.", bold=True, size=BODY)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_run(p, "Respected Sir/Ma\u2019am,", size=BODY)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "With reference to the above subject, we wish to re-appoint you as the Statutory Auditors of our company.", size=BODY)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    add_run(p, f"Please let us know that whether you are duly qualified and eligible for this re-appointment as per the provisions of Section 139 and 141 of the Companies Act 2013 and rules made there under and provide your written consent to act as Statutory Auditors of the Company from the conclusion of this Annual General Meeting of the company till the Annual General Meeting held for Financial Year ending {aud['tenure_end_fy']} for term of {aud['term_years']} Years.", size=BODY)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "You are also requested to issue a certificate to the Company under Section 139(1) of the Companies Act 2013.", size=BODY)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    add_run(p, "Thanking you.", size=BODY)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_run(p, "Yours faithfully,", size=BODY)

    # Compact signing block \u2014 kept together so it can't orphan.
    signer = signing["first"]
    def _kw(p):
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True
        return _tight(p)
    p = _kw(doc.add_paragraph()); add_run(p, f"For {company['name'].upper()}", bold=True, size=BODY)
    p = _kw(doc.add_paragraph()); add_run(p, "", size=BODY)
    p = _kw(doc.add_paragraph()); add_run(p, "", size=BODY)
    p = _kw(doc.add_paragraph()); add_run(p, signer["name"].upper(), bold=True, size=BODY)
    p = _kw(doc.add_paragraph()); add_run(p, signer.get("designation", "Director"), bold=True, size=BODY)
    p = _tight(doc.add_paragraph())
    p.paragraph_format.keep_together = True
    add_run(p, f"DIN: {signer['din']}", bold=True, size=BODY)


# =====================================================================
# SECTION 12: CONSENT LETTER FROM AUDITOR (auditor's letterhead)
# =====================================================================
def section_consent_letter(doc, data):
    if not yes(data["toggles"], "AuditorReappointment"):
        return
    aud = data["auditor"]
    company = data["company"]
    agm = data["agm"]

    # Term FY range — derived backward from Tenure End FY (year of last AGM in tenure).
    # User enters only the end year; start is auto-computed using Term Length.
    try:
        years_in_end = re.findall(r"\d{4}", str(aud.get("tenure_end_fy", "")))
        if not years_in_end:
            raise ValueError("no year in tenure_end_fy")
        end_year = int(years_in_end[-1])
        term = int(str(aud.get("term_years", "5")).strip() or "5")
        first_fy_start = end_year - term                                       # 2025
        term_short = f"{first_fy_start}-{str(first_fy_start + 1)[-2:]}"        # "2025-26"
        term_long_end = f"{end_year - 1}-{end_year}"                           # "2029-2030"
    except (ValueError, TypeError):
        term_short = "(Start FY)"
        term_long_end = "(End FY)"

    # Compact — 10pt body, tight spacing so the whole letter fits on one page.
    BODY = 10
    def _tight(p):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        return p

    # Letterhead placeholder — auditor's own letterhead (not company)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Letterhead of auditor", italic=True, bold=True, size=BODY)

    add_centered_heading(doc, "CONSENT LETTER CUM ELIGIBILITY CERTIFICATE", size=12, underline=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(4)
    add_run(p, f"Date: {agm.get('notice_date', '')}", bold=True, size=BODY)

    p = _tight(doc.add_paragraph()); add_run(p, "To,", size=BODY)
    p = _tight(doc.add_paragraph()); add_run(p, "The Board of Directors,", size=BODY)
    p = _tight(doc.add_paragraph()); add_run(p, company["name"], bold=True, size=BODY)
    p = _tight(doc.add_paragraph())
    p.paragraph_format.space_after = Pt(6)
    add_run(p, f"Address: {company['address']}", bold=True, size=BODY)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    add_run(p, f"Sub: Re-Appointment as Statutory Auditors of the Company for Financial Year {term_short} to Financial Year {term_long_end}.", bold=True, underline=True, size=BODY)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "We thank you for your letter, informing us that it is proposed to Re-appoint ", size=BODY)
    add_run(p, f"{aud['firm_name']}", bold=True, size=BODY)
    add_run(p, f". as Statutory Auditors of the Company for Financial Year {term_short} to Financial Year {term_long_end}.", size=BODY)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "We would be happy to accept the appointment, if made.", size=BODY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Further, as required under section 139(1) of the Companies Act, 2013 (“the Act”), we further certify, declare and confirm that:", size=BODY)

    declarations = [
        "The firm is eligible for re-appointment and is not disqualified for re-appointment under the Act, the Chartered Accountants Act, 1949 and the rules or regulations made there under;",
        "The proposed re-appointment is as per the term provided under the Act;",
        "The proposed re-appointment is within the limits laid down by or under the authority of the Act;",
        "No proceedings are pending against the firm or any of its partners with respect to professional matters of conduct;",
        "The firm satisfies the criteria provided in section 141 of the Act and is eligible to re-appointment.",
    ]
    for i, decl in enumerate(declarations, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # left_indent removed — keep resolution body aligned with item text above
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f"{i}) ", size=BODY)
        add_run(p, decl, size=BODY)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Assuring you of our best attention always. We further assure you our best professional services at all times.", size=BODY)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    add_run(p, "Thanking You,", size=BODY)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_run(p, "Yours Faithfully,", size=BODY)

    # Firm + partner signature block — keep together so the partner name
    # never orphans onto a near-empty next page.
    def _keep(p):
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True
        return p

    p = _keep(doc.add_paragraph()); _tight(p); add_run(p, aud["firm_name"], bold=True, size=BODY)
    p = _keep(doc.add_paragraph()); _tight(p); add_run(p, aud["designation"], bold=True, size=BODY)
    p = _keep(doc.add_paragraph()); _tight(p); add_run(p, f"FRN: {aud['frn']}", bold=True, size=BODY)

    # One small visual gap (also kept-with-next so it travels with the block).
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    _keep(spacer)

    p = _keep(doc.add_paragraph()); _tight(p); add_run(p, f"({aud.get('partner_name', '')})", bold=True, size=BODY)
    p = _keep(doc.add_paragraph()); _tight(p); add_run(p, "PARTNER/PROPRIETOR", bold=True, size=BODY)
    p = doc.add_paragraph(); _tight(p)
    p.paragraph_format.keep_together = True
    add_run(p, f"M. NO. {aud.get('partner_mno', '')}", bold=True, size=BODY)


# =====================================================================
# SECTION 13: DETAILS OF DESIGNATED PERSON (RoC letter — company letterhead)
# =====================================================================
def section_designated_person_letter(doc, data):
    company = data["company"]
    agm = data["agm"]
    signing = data["signing"]

    add_centered_heading(doc, "DETAILS OF DESIGNATED PERSON", size=12, underline=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, f"Date: {agm.get('notice_date', '')}", bold=True)

    add_para(doc, "")
    # Tight address block — no gap between To / RoC / MCA / RoC address lines.
    def _tight(p):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        return p

    p = _tight(doc.add_paragraph()); add_run(p, "To,")
    p = _tight(doc.add_paragraph()); add_run(p, "The Registrar of Companies")
    p = _tight(doc.add_paragraph()); add_run(p, "Ministry of Corporate Affairs,")

    roc = (company.get("roc_address") or "").strip()
    if roc:
        for line in roc.split("\n"):
            line = line.strip()
            if line:
                p = _tight(doc.add_paragraph()); add_run(p, line)

    add_para(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "Subject: ", bold=True, underline=True)
    add_run(p, "Details of designated person for furnishing and extending co-operation for providing information to the Registrar or any other authorized officer with respect to beneficial interest in shares of the Company.", bold=True, underline=True)

    add_para(doc, "")
    add_para(doc, "Respected Sir/ Ma’am,")

    add_para(doc, "")
    _justify(doc, "With reference to the cited subject and pursuant to the compliance of Rule 9(7) of Companies (Management and Administration) Rules, 2014 inserted by MCA Notification dated October 27, 2023, we hereby providing you the details of person, who shall be responsible for furnishing and extending co-operation for providing information to the Registrar or any other authorized officer with respect to beneficial interest in shares of the Company.")

    add_para(doc, "")
    _justify(doc, "The details of the said designated person are as follows:")

    dp = signing.get("designated_person", {})
    add_bordered_table(doc,
        ["S. No.", "Name of the Designated Person", "Designation", "DIN"],
        [["1", dp.get("name", ""), "Director", dp.get("din", "")]],
        col_widths_inches=[0.7, 2.8, 1.5, 1.5])

    add_para(doc, "")
    _justify(doc, "You are requested to kindly take the same in your record.")

    add_para(doc, "")
    p = doc.add_paragraph()
    add_run(p, f"For {company['name'].upper()}", bold=True)

    # Signature stamp space
    add_para(doc, "")
    add_para(doc, "")
    add_para(doc, "")

    # Two signers side-by-side, left-aligned (matches the screenshot layout)
    tbl = doc.add_table(rows=3, cols=2)
    tbl.autofit = True
    first = signing["first"]
    second = signing["second"]
    rows = [
        ("name", lambda s: s["name"].upper()),
        ("designation", lambda s: s.get("designation", "Director")),
        ("din", lambda s: f"(DIN: {s['din']})"),
    ]
    for r_idx, (key, fmt) in enumerate(rows):
        for col_idx, signer in enumerate([first, second]):
            cell = tbl.rows[r_idx].cells[col_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, fmt(signer), bold=True)
            _set_cell_borders(cell, "nil")


# =====================================================================
# MASTER ORCHESTRATOR
# =====================================================================
def build_final_set(data, out_path):
    doc = new_doc()
    add_letterhead(doc, data)

    # 1. Notice (multi-page with its own letterheads & page breaks)
    section_notice(doc, data)

    # 2. Director's Report
    _new_page_with_letterhead(doc, data)
    section_directors_report(doc, data)

    # 3. AOC-2 (Annexure A)
    _new_page_with_letterhead(doc, data)
    section_aoc2(doc, data)

    # 4. Equity List of Shareholders (LANDSCAPE)
    _begin_landscape_page(doc, data)
    section_shareholders_list(doc, data)

    # 5. Attendance of AGM (back to PORTRAIT)
    _begin_portrait_page(doc, data)
    section_attendance_agm(doc, data)

    # 6. List of Directors
    _new_page_with_letterhead(doc, data)
    section_directors_list(doc, data)

    # 7. Director Attendance Sheet (LANDSCAPE)
    _begin_landscape_page(doc, data)
    section_director_attendance(doc, data)

    # 8. Resolution - Director Regularization (back to PORTRAIT)
    if yes(data["toggles"], "RegulariseAdditionalDirector") and data.get("regularizations"):
        _begin_portrait_page(doc, data)
        section_resolution_director(doc, data)
    else:
        _begin_portrait_page(doc, data)

    # 9. Resolution - RPT
    if yes(data["toggles"], "RPT_Required") and data.get("rpts"):
        _new_page_with_letterhead(doc, data)
        section_resolution_rpt(doc, data)

    # 10. Resolution - Auditor
    if yes(data["toggles"], "AuditorReappointment"):
        _new_page_with_letterhead(doc, data)
        section_resolution_auditor(doc, data)

    # 11. Intimation Letter
    if yes(data["toggles"], "AuditorReappointment"):
        _new_page_with_letterhead(doc, data)
        section_intimation_letter(doc, data)

    # 12. Consent Letter from Auditor (auditor's letterhead — no company letterhead)
    if yes(data["toggles"], "AuditorReappointment"):
        doc.add_page_break()
        section_consent_letter(doc, data)

    # 13. Details of Designated Person (RoC letter, company letterhead)
    # Driven by the explicit FirstDesignatedPersonDeclaration toggle —
    # YES = first Rule 9(7) filing, include the letter; NO = already on RoC record.
    if yes(data["toggles"], "FirstDesignatedPersonDeclaration"):
        _new_page_with_letterhead(doc, data)
        section_designated_person_letter(doc, data)

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    HERE = os.path.dirname(os.path.abspath(__file__))
    PROJECT_ROOT = os.path.dirname(HERE)
    data = load_master(os.path.join(PROJECT_ROOT, "master", "Master_Input.xlsx"))
    build_final_set(data, os.path.join(PROJECT_ROOT, "output", "Final_Set.docx"))
