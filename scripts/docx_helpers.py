"""
docx_helpers.py - Shared building blocks used by every generator.

The most-repeated structures in your AGM document set are:
  * Letterhead (Company name centered, CIN, address, email/phone in a row)
  * Signing block at the end (one or two directors, with DIN)
  * Page setup (A4 portrait, 1" margins, Arial)
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# Document setup
# ============================================================
def new_doc():
    """Return a fresh Document with default font Cambria 11 / 1.35 line spacing, A4 portrait, 1" margins."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Cambria"
    style.font.size = Pt(11)
    # Comfortable reading rhythm: 1.35 line spacing gives each line breathing
    # room without ballooning the page count, and 6pt after each paragraph
    # creates a clear visual break between sentences/clauses. Heading
    # space_before stacks on top so section breaks still feel distinct.
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.35
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Cambria")
    rfonts.set(qn("w:hAnsi"), "Cambria")

    for section in doc.sections:
        # A4 portrait dimensions — default python-docx is US Letter
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    return doc


# ============================================================
# Run / paragraph helpers
# ============================================================
def add_run(p, text, bold=False, italic=False, underline=False, size=None, font="Cambria"):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.underline = underline
    r.font.name = font
    if size:
        r.font.size = Pt(size)
    return r


def add_para(doc, text="", bold=False, align=None, size=None, after=None, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size)
    # Default body spacing — overridable via `after` (use after=0 for true tight blocks).
    p.paragraph_format.space_after = Pt(after if after is not None else 6)
    return p


def add_justified(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text)
    return p


# ============================================================
# Letterhead - placed at top of every document
# ============================================================
def add_letterhead(doc, data):
    """Header strip: Company name (bold center), CIN, address, contact line, hr."""
    c = data["company"]

    def _tight(p):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15

    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, c["name"], bold=True, size=18)
    _tight(p)
    if c.get("old_name"):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2, f"(Formerly Known As {c['old_name']})", italic=True, size=10)
        _tight(p2)

    # CIN
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"CIN: {c['cin']}", bold=True, size=11)
    _tight(p)

    # Address
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"Reg. office: {c['address']}", bold=True, size=10)
    _tight(p)

    # Email + phone (with website if set)
    contact_parts = [f"E-mail: {c['email']}"]
    if c.get("website"):
        contact_parts.append(f"Website: {c['website']}")
    contact_parts.append(f"Contact No. {c['phone']}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "      ".join(contact_parts), bold=True, size=10)
    _tight(p)

    # Horizontal line below header
    add_horizontal_rule(doc)


def add_horizontal_rule(doc):
    """Insert a horizontal black line as a paragraph border."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============================================================
# Signing block at end of document
# ============================================================
def add_single_signer_block(doc, signer, date_str, place_str,
                             closing="By Order of the Board of Directors",
                             company_name=None, reg_office_address=None):
    """
    Single signer (used in Notice). Centered-right.

    If reg_office_address is provided, lays out as a 2-column borderless table:
      Left  = Registered Office + address, then Date/Place at bottom
      Right = closing line + FOR <COMPANY> + signer name/designation/DIN
    """
    if reg_office_address:
        add_para(doc, "")
        tbl = doc.add_table(rows=1, cols=2)
        tbl.autofit = False
        left, right = tbl.rows[0].cells
        left.width = Inches(3.6)
        right.width = Inches(3.6)
        _set_cell_borders(left, "nil")
        _set_cell_borders(right, "nil")

        # LEFT cell: Registered Office + Date + Place
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, "Registered Office:", bold=True)
        p = left.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, reg_office_address, bold=True)
        left.add_paragraph()
        left.add_paragraph()
        left.add_paragraph()
        left.add_paragraph()
        if date_str:
            p = left.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, f"Date- {date_str}", bold=True)
        if place_str:
            p = left.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, f"Place- {place_str}", bold=True)

        # RIGHT cell: closing line + FOR COMPANY + signer
        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, closing, bold=True)
        if company_name:
            right.add_paragraph()
            p = right.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, f"FOR {company_name.upper()}", bold=True)
        # Spacer for signature
        right.add_paragraph()
        right.add_paragraph()
        right.add_paragraph()
        p = right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, signer["name"].upper(), bold=True)
        p = right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, signer.get("designation", "Director"), bold=True)
        p = right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, f"DIN: {signer['din']}", bold=True)
        return

    add_para(doc, "")
    p = doc.add_paragraph()
    add_run(p, closing, bold=True)

    if company_name:
        add_para(doc, "")
        p = doc.add_paragraph()
        add_run(p, f"FOR {company_name.upper()}", bold=True)

    # Spacer for signature
    add_para(doc, "")
    add_para(doc, "")
    add_para(doc, "")

    p = doc.add_paragraph()
    add_run(p, signer["name"].upper(), bold=True)
    p = doc.add_paragraph()
    add_run(p, signer.get("designation", "Director"), bold=True)
    p = doc.add_paragraph()
    add_run(p, f"DIN: {signer['din']}", bold=True)

    add_para(doc, "")
    if date_str:
        p = doc.add_paragraph()
        add_run(p, f"Date: {date_str}", bold=True)
    if place_str:
        p = doc.add_paragraph()
        add_run(p, f"Place: {place_str}", bold=True)


def add_two_signer_block(doc, first, second, date_str, place_str, company_name):
    """Two signers side-by-side (used in Board Report, Resolutions)."""
    add_para(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, "For and on behalf of the Board of Directors", bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, f"FOR {company_name.upper()}", bold=True)

    # Spacer
    add_para(doc, "")
    add_para(doc, "")
    add_para(doc, "")

    # 1x2 table for two signers
    tbl = doc.add_table(rows=3, cols=2)
    tbl.autofit = True

    cells_top = tbl.rows[0].cells
    cells_mid = tbl.rows[1].cells
    cells_bot = tbl.rows[2].cells

    for cells, key in [(cells_top, "name"), (cells_mid, "designation"), (cells_bot, "din")]:
        for i, signer in enumerate([first, second]):
            cell = cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if key == "name":
                add_run(p, signer["name"].upper(), bold=True)
            elif key == "designation":
                add_run(p, signer.get("designation", "Director"), bold=True)
            else:
                add_run(p, f"(DIN: {signer['din']})", bold=True)
            # remove cell borders for clean look
            _set_cell_borders(cell, "nil")

    add_para(doc, "")
    if date_str:
        p = doc.add_paragraph()
        add_run(p, f"Date: {date_str}", bold=True)
    if place_str:
        p = doc.add_paragraph()
        add_run(p, f"Place: {place_str}", bold=True)


# ============================================================
# Table page-break helpers
# ============================================================
# Threshold: tables with this many data rows or fewer try to stay on one page.
# Bigger tables (e.g. landscape Shareholders/Director-Attendance lists) split
# row-by-row instead so they don't blow up.
_KEEP_TOGETHER_MAX_ROWS = 8


def _row_cant_split(row):
    """Prevent a single row from being split horizontally between pages."""
    trPr = row._tr.get_or_add_trPr()
    # Remove any existing cantSplit before re-adding (idempotent).
    for existing in trPr.findall(qn("w:cantSplit")):
        trPr.remove(existing)
    cant = OxmlElement("w:cantSplit")
    cant.set(qn("w:val"), "true")
    trPr.append(cant)


def _row_keep_with_next(row):
    """Bind this row to the next so they don't separate across pages.
    Implemented by setting keep_with_next on every paragraph in every cell."""
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = True


def _apply_table_pagination(tbl):
    """Standard table-pagination policy:
      - every row: cantSplit (no mid-row breaks)
      - if total rows <= threshold: bind every row to the next so the whole
        table travels as a single unit.
    """
    rows = list(tbl.rows)
    for row in rows:
        _row_cant_split(row)
    if len(rows) <= _KEEP_TOGETHER_MAX_ROWS:
        for row in rows[:-1]:
            _row_keep_with_next(row)


def _set_cell_borders(cell, border_val="single"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), border_val)
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tcBorders.append(b)
    tc_pr.append(tcBorders)


# ============================================================
# Tables - bordered table builder for board reports
# ============================================================
def add_bordered_table(doc, headers, rows, header_bold=True, col_widths_inches=None):
    """
    Build a table with all-borders, header row in bold.
    rows = list of lists (one inner list per row).
    """
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"

    # Header
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, str(h), bold=header_bold)
        _shade_cell(cell, "DCE6F2")

    # Body
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            add_run(p, str(val) if val is not None else "")

    if col_widths_inches:
        for i, w in enumerate(col_widths_inches):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    _apply_table_pagination(tbl)
    return tbl


def add_aoc2_block_table(doc, block_label, lettered_rows,
                         col_widths_inches=(0.5, 4.0, 2.0)):
    """Build an AOC-2 "Block-N" framed table.

    `block_label`     — the merged header text (e.g. "Block-1").
    `lettered_rows`   — list of (letter, description, value) tuples;
                        e.g. ("a)", "Corporate identity number...", "AFGFS5599P").
    Layout: 3-column table; row 0 is the merged Block label spanning all columns.
    """
    n_cols = 3
    tbl = doc.add_table(rows=1 + len(lettered_rows), cols=n_cols)
    tbl.style = "Table Grid"

    header = tbl.rows[0].cells
    merged = header[0]
    merged.merge(header[1]).merge(header[2])
    merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = merged.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, block_label, bold=True)
    _shade_cell(merged, "DCE6F2")

    for r_idx, (letter, desc, value) in enumerate(lettered_rows, start=1):
        c1, c2, c3 = tbl.rows[r_idx].cells
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, str(letter))
        p = c2.paragraphs[0]
        add_run(p, str(desc))
        p = c3.paragraphs[0]
        add_run(p, str(value) if value is not None else "")

    for i, w in enumerate(col_widths_inches):
        for row in tbl.rows:
            row.cells[i].width = Inches(w)
    _apply_table_pagination(tbl)
    return tbl


def add_aoc2_count_box(doc, label_text, count_value, label_width_inches=5.5, box_width_inches=0.8):
    """Render a one-row 2-cell table inside a single bordered box: label on the
    left, count on the right, separated by a vertical divider.

    Mirrors the PDF AOC-2 "Number of contracts... ⎕" layout.
    """
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    label_cell, box_cell = tbl.rows[0].cells
    label_cell.width = Inches(label_width_inches)
    box_cell.width = Inches(box_width_inches)
    label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    box_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = label_cell.paragraphs[0]
    add_run(p, label_text)
    p = box_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, str(count_value), bold=True)
    _apply_table_pagination(tbl)
    return tbl


def add_keyvalue_table(doc, rows, label_width_inches=2.0, value_width_inches=4.5,
                       label_bold=True):
    """Build a 2-column non-header table for label/value pairs.

    Matches the Reserves / Dividend style in the gold-standard Director's Report:
    each row has a bold left label and a free-form right value.
    """
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    for r_idx, (label, value) in enumerate(rows):
        lc = tbl.rows[r_idx].cells[0]
        vc = tbl.rows[r_idx].cells[1]
        lc.width = Inches(label_width_inches)
        vc.width = Inches(value_width_inches)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = lc.paragraphs[0]
        add_run(p, str(label), bold=label_bold)
        p = vc.paragraphs[0]
        add_run(p, str(value) if value is not None else "")
    _apply_table_pagination(tbl)
    return tbl


def _shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


# ============================================================
# Misc
# ============================================================
def add_centered_heading(doc, text, size=13, underline=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    add_run(p, text, bold=True, size=size, underline=underline)
    return p


def add_section_heading(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    add_run(p, text, bold=True, size=size, underline=True)
    return p


def add_numbered_item(doc, number, text_bold, text_rest=""):
    """Like '1. **Title:** rest of text' as one paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    add_run(p, f"{number}. ", bold=True)
    add_run(p, text_bold, bold=True)
    if text_rest:
        add_run(p, text_rest)
    return p
